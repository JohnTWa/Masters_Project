import difflib
from docx import Document
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def generate_html_diff(file1_path, file2_path, output_html_path):
    """
    Generates an HTML file showing a single merged text where:
      - Removed (or replaced) characters are highlighted in red with strikethrough.
      - Added (or replacing) characters are highlighted in green with underline.
    Assumes that each line in file1 corresponds to the same line in file2.
    """
    # Read the contents of both files as lists of lines.
    with open(file1_path, 'r') as f1, open(file2_path, 'r') as f2:
        original_lines = f1.readlines()
        modified_lines = f2.readlines()
    
    merged_lines = []
    
    # Process each corresponding pair of lines.
    for orig, mod in zip(original_lines, modified_lines):
        sm = difflib.SequenceMatcher(None, orig, mod)
        merged_line_parts = []
        
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                merged_line_parts.append(orig[i1:i2])
            elif tag == 'delete':
                removed_text = orig[i1:i2]
                merged_line_parts.append(f'<span class="delete">{removed_text}</span>')
            elif tag == 'insert':
                inserted_text = mod[j1:j2]
                merged_line_parts.append(f'<span class="insert">{inserted_text}</span>')
            elif tag == 'replace':
                removed_text = orig[i1:i2]
                inserted_text = mod[j1:j2]
                merged_line_parts.append(f'<span class="delete">{removed_text}</span>')
                merged_line_parts.append(f'<span class="insert">{inserted_text}</span>')
                
        # Join the parts and replace newline characters with <br>.
        merged_line = ''.join(merged_line_parts).replace('\n', '<br>')
        merged_lines.append(merged_line)
    
    # Construct the HTML content.
    html_content = [
        '<html>',
        '<head>',
        '<meta charset="utf-8">',
        '<style>',
        'body { font-family: monospace; white-space: pre-wrap; }',
        '.delete { color: red; text-decoration: line-through; }',
        '.insert { color: green; text-decoration: underline; }',
        '</style>',
        '</head>',
        '<body>',
        '\n'.join(merged_lines),
        '</body>',
        '</html>'
    ]
    
    # Write the HTML content to the output file.
    with open(output_html_path, 'w') as out_file:
        out_file.write('\n'.join(html_content))

def add_strikethrough(run):
    """
    Adds strikethrough formatting to a run in a DOCX document.
    """
    r = run._element
    rPr = r.get_or_add_rPr()
    strike = OxmlElement('w:strike')
    strike.set(qn('w:val'), "true")
    rPr.append(strike)

def generate_docx_diff(file1_path, file2_path, output_docx_path):
    """
    Generates a DOCX file showing a single merged text where:
      - Removed (or replaced) characters are shown in red with strikethrough.
      - Added (or replacing) characters are shown in green with underline.
    
    Assumes each line in file1 corresponds to the same line in file2.
    Instead of creating a new paragraph for each line, it adds a soft newline (line break)
    at the end of each line without extra spacing.
    """
    # Read the contents of both files, stripping newlines.
    with open(file1_path, 'r') as f1, open(file2_path, 'r') as f2:
        original_lines = f1.read().splitlines()
        modified_lines = f2.read().splitlines()
    
    document = Document()
    # Create a single paragraph to hold the merged text.
    paragraph = document.add_paragraph()
    
    # Process each corresponding pair of lines.
    for idx, (orig, mod) in enumerate(zip(original_lines, modified_lines)):
        sm = difflib.SequenceMatcher(None, orig, mod)
        
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'equal':
                paragraph.add_run(orig[i1:i2])
            elif tag == 'delete':
                run = paragraph.add_run(orig[i1:i2])
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # red
                add_strikethrough(run)
            elif tag == 'insert':
                run = paragraph.add_run(mod[j1:j2])
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # green
                run.font.underline = True
            elif tag == 'replace':
                # Show the removed part from the original.
                run = paragraph.add_run(orig[i1:i2])
                run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # red
                add_strikethrough(run)
                # Then show the inserted part from the modified.
                run = paragraph.add_run(mod[j1:j2])
                run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # green
                run.font.underline = True

        # Add a soft line break (without extra paragraph spacing) after each line, except after the last.
        if idx < len(original_lines) - 1:
            paragraph.add_run().add_break()
    
    # Save the DOCX document.
    document.save(output_docx_path)

if __name__ == '__main__':
    generate_html_diff('files/text_files/t1_transmission_text.txt', 'files/text_files/t2_received_text.txt', 'files/text_files/t3_text_diff.html')
    generate_docx_diff('files/text_files/t1_transmission_text.txt', 'files/text_files/t2_received_text.txt', 'files/text_files/t3_text_diff.docx')